import os
import re
import json
import gradio as gr
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from dotenv import load_dotenv
import spacy
import nltk
from typing import List, Tuple, Optional, Callable, Dict, Any
import uuid
import time
import hashlib
import threading
from concurrent.futures import ThreadPoolExecutor

# File readers
import PyPDF2
import docx                        # python-docx  → Word files
import openpyxl                    # openpyxl     → Excel files
from PIL import Image              # Pillow       → Images
import pytesseract                 # pytesseract  → OCR on images
load_dotenv()
openai_api_key   = os.getenv("OPENAI_API_KEY")
OPENAI_LLM_MODEL = os.getenv("OPENAI_LLM_MODEL", "gpt-4o-mini")  # override via OPENAI_LLM_MODEL in .env

# ── User directory ────────────────────────────────────────────────────────────
users = {
    "john.doe123@example.com":   {"username": "User1", "id": str(uuid.uuid4()), "org": "Org1"},
    "jane.smith456@example.com": {"username": "User2", "id": str(uuid.uuid4()), "org": "Org1"},
    "bob.jones789@example.com":  {"username": "User3", "id": str(uuid.uuid4()), "org": "Org2"},
    "alice.brown321@example.com":{"username": "User4", "id": str(uuid.uuid4()), "org": "Org2"},
    "mike.wilson654@example.com":{"username": "User5", "id": str(uuid.uuid4()), "org": "Org2"},
    "admin@example.com":         {"username": "Admin", "id": str(uuid.uuid4()), "org": None},
}

# ── NLP setup ─────────────────────────────────────────────────────────────────
try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    nltk.download("punkt")
try:
    spacy_nlp = spacy.load("en_core_web_sm")
except OSError:
    spacy.cli.download("en_core_web_sm")
    spacy_nlp = spacy.load("en_core_web_sm")

# spaCy refuses to parse text longer than `max_length` (default 1_000_000).
# Large PDFs blow past this. We still window the text below, but this gives
# headroom for borderline files and prevents a hard crash mid-upload.
spacy_nlp.max_length = 5_000_000

# How much text we feed spaCy in one go. Stays comfortably under the limit
# above and keeps peak RAM reasonable (spaCy uses ~10MB per 100K chars).
SPACY_WINDOW_CHARS = 800_000

# Max concurrent OpenAI embedding requests. Tuned to be aggressive but stay
# well under Tier-1 / Tier-2 rate limits for text-embedding-3-small.
EMBED_WORKERS = 8

# Hard cap on uploaded file size. Enforced both by Gradio (`max_file_size` on
# launch) and by an explicit server-side check inside `process_file`.
MAX_UPLOAD_MB = 50
MAX_UPLOAD_BYTES = MAX_UPLOAD_MB * 1024 * 1024


# ══════════════════════════════════════════════════════════════════════════════
#  FILE-TYPE AGENT
#  Decides how to read a file and returns plain text.
# ══════════════════════════════════════════════════════════════════════════════

# Tool definitions the LLM can choose from
FILE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Extract text from a PDF file using PyPDF2.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the PDF file"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_word",
            "description": "Extract text from a Microsoft Word (.docx) file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the .docx file"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_excel",
            "description": "Extract all cell values from a Microsoft Excel (.xlsx / .xls) file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the Excel file"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_image",
            "description": "Use OCR (pytesseract) to extract text from an image file (png, jpg, jpeg, tiff, bmp, webp).",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the image file"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_text",
            "description": "Read a plain-text file (.txt, .md, .csv, .json, .xml, .html, etc.).",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the text file"}
                },
                "required": ["file_path"]
            }
        }
    },
]


# ── Concrete tool implementations ─────────────────────────────────────────────

def read_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text


def _read_pdf_with_pages(file_path: str) -> List[Dict[str, Any]]:
    """Page-aware variant of `read_pdf`. Returns a list of segments
    [{'text': <page text>, 'page': <1-indexed page number>}, ...].

    Used by `file_type_agent` when the LLM picks `read_pdf`, so that
    `process_document` can attach per-chunk page numbers to the Qdrant
    payload. Pages whose extracted text is empty are skipped to avoid
    polluting the chunker with blank segments."""
    segments: List[Dict[str, Any]] = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                segments.append({"text": page_text, "page": page_num})
    return segments


def read_word(file_path: str) -> str:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def read_excel(file_path: str) -> str:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    lines = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        lines.append(f"=== Sheet: {sheet} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)


def read_image(file_path: str) -> str:
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
    return text


def read_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# Map tool names → functions
TOOL_FUNCTIONS = {
    "read_pdf":   read_pdf,
    "read_word":  read_word,
    "read_excel": read_excel,
    "read_image": read_image,
    "read_text":  read_text,
}


def file_type_agent(
    file_path: str, filename: str, llm_client
) -> Tuple[str, str, List[Dict[str, Any]]]:
    """
    An agent that decides which reader tool to use based on the filename,
    calls it, and returns `(extracted_text, agent_reasoning_log, segments)`.

    `segments` is a list of `{'text': str, 'page': Optional[int]}` dicts:
      - For PDFs, one segment per page so downstream chunking can attach
        page numbers to each Qdrant point's payload.
      - For every other file type, a single segment with `page=None`
        containing the full extracted text (i.e., behavior unchanged for
        chunking; page citations just won't appear in the chatbot's
        Sources line).

    The LLM chooses the right tool — this is the core agentic pattern:
    Observe → Decide → Act → Return result
    """
    log_lines = []

    # ── Step 1: Ask the LLM to pick a tool ───────────────────────────────────
    system_prompt = (
        "You are a file-reading agent. Given a filename you must call exactly ONE "
        "of the available tools to extract its text content. "
        "Choose the most appropriate tool based on the file extension. "
        "Do not explain yourself — just call the tool."
    )
    user_message = (
        f"Extract text from this file: '{filename}'\n"
        f"Full path: {file_path}"
    )

    log_lines.append(f"🤖 Agent received file: {filename}")
    log_lines.append(f"🔍 Asking LLM to pick the right reader tool...")

    response = llm_client.chat.completions.create(
        model=OPENAI_LLM_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_message},
        ],
        tools=FILE_TOOLS,
        tool_choice="required",   # force the LLM to always call a tool
    )

    message = response.choices[0].message

    # ── Step 2: Execute the chosen tool ──────────────────────────────────────
    if not message.tool_calls:
        log_lines.append("⚠️  LLM returned no tool call. Falling back to plain-text reader.")
        tool_name = "read_text"
        tool_args = {"file_path": file_path}
    else:
        tool_call = message.tool_calls[0]
        tool_name = tool_call.function.name
        tool_args = json.loads(tool_call.function.arguments)
        tool_args["file_path"] = file_path   # always use the real server path

    log_lines.append(f"🛠️  Agent chose tool: **{tool_name}**")

    if tool_name not in TOOL_FUNCTIONS:
        raise ValueError(f"Unknown tool chosen by agent: {tool_name}")

    log_lines.append(f"⚙️  Executing {tool_name}...")

    # For PDFs we use the page-aware extractor so that per-chunk page
    # numbers can flow all the way into Qdrant. Every other tool still
    # returns plain text, which we wrap into a single page-less segment
    # so downstream code only has to handle one shape.
    segments: List[Dict[str, Any]]
    if tool_name == "read_pdf":
        segments = _read_pdf_with_pages(file_path)
        extracted_text = "\n\n".join(s["text"] for s in segments)
        log_lines.append(f"📄 Extracted {len(segments)} page(s) with text.")
    else:
        extracted_text = TOOL_FUNCTIONS[tool_name](file_path)
        segments = [{"text": extracted_text, "page": None}]

    char_count = len(extracted_text)
    log_lines.append(f"✅ Extraction complete — {char_count:,} characters extracted.")

    if char_count < 50:
        log_lines.append("⚠️  Very little text extracted. File may be empty or unsupported.")

    agent_log = "\n".join(log_lines)
    return extracted_text, agent_log, segments


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT PROCESSOR  (unchanged logic, same class as before)
# ══════════════════════════════════════════════════════════════════════════════

class DocumentProcessor:
    def __init__(self,
                 collection_name="aiml_vector_db",
                 qdrant_host="localhost",
                 qdrant_port=6333,
                 chunk_strategy="semantic",
                 chunk_size=800,
                 chunk_overlap=150,
                 embedding_model="text-embedding-3-small",
                 batch_size=32):
        self.client = QdrantClient(host=qdrant_host, port=qdrant_port)
        self.collection_name = collection_name
        self.next_id = 0
        # Guards all reads/writes of self.next_id so two concurrent uploads
        # (Gradio queue concurrency > 1) can't reserve overlapping ID ranges.
        self._id_lock = threading.Lock()
        self._create_collection(vector_size=1536, distance=Distance.COSINE)
        self.chunk_strategy = chunk_strategy.lower()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        if self.chunk_strategy == "recursive_char":
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ".", " ", ""]
            ).split_text
        elif self.chunk_strategy == "semantic":
            self.splitter = lambda text: self._semantic_split(text)
        else:
            raise ValueError(f"Unsupported chunking strategy: {chunk_strategy}")
        self.embeddings = OpenAIEmbeddings(
            model=embedding_model,
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            timeout=60.0,
            max_retries=3,
        )
        self.batch_size = batch_size
        self.processed_files = self._fetch_processed_files()

    def _create_collection(self, vector_size: int, distance=Distance.COSINE):
        collections = self.client.get_collections().collections
        if not any(col.name == self.collection_name for col in collections):
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(size=vector_size, distance=distance)
            )
            print(f"Created collection: {self.collection_name}")
        else:
            print(f"Collection {self.collection_name} already exists")

    def _fetch_processed_files(self) -> dict:
        """Scroll the collection to (a) rebuild the source_file → user_email
        map and (b) advance self.next_id past the largest existing integer
        point ID. Without (b), restarting the process would reset next_id to
        0 and the next upload would silently overwrite the first N rows."""
        try:
            processed_files = {}
            max_int_id = -1
            scroll_result = self.client.scroll(
                collection_name=self.collection_name,
                limit=256,
                with_payload=True,
                scroll_filter=None
            )
            points, next_offset = scroll_result
            while points:
                for point in points:
                    source_file = point.payload.get("source_file")
                    user_email  = point.payload.get("user_email")
                    if source_file and user_email:
                        processed_files[source_file] = user_email
                    # Only int IDs matter for the counter; UUID-string IDs
                    # (if any future code uses them) are ignored.
                    if isinstance(point.id, int) and point.id > max_int_id:
                        max_int_id = point.id
                if next_offset is None:
                    break
                scroll_result = self.client.scroll(
                    collection_name=self.collection_name,
                    limit=256,
                    offset=next_offset,
                    with_payload=True,
                    scroll_filter=None
                )
                points, next_offset = scroll_result
            # max() guards against an in-flight upload that has reserved IDs
            # above what's currently visible in the scroll snapshot.
            with self._id_lock:
                self.next_id = max(self.next_id, max_int_id + 1)
            return processed_files
        except Exception as e:
            print(f"Error fetching processed files: {str(e)}")
            return {}

    def _iter_spacy_windows(self, text: str):
        """Yield ~SPACY_WINDOW_CHARS-sized slices of `text`, snapping to the
        nearest paragraph break to avoid splitting a sentence across windows.
        Keeps spaCy's memory bounded and dodges the 1M-char default ceiling."""
        if len(text) <= SPACY_WINDOW_CHARS:
            yield text
            return
        start = 0
        n = len(text)
        while start < n:
            end = min(start + SPACY_WINDOW_CHARS, n)
            if end < n:
                snap = text.rfind("\n\n", start, end)
                if snap > start + SPACY_WINDOW_CHARS // 2:
                    end = snap
            yield text[start:end]
            start = end

    def _semantic_split(self, text: str) -> List[str]:
        chunks, current_chunk, current_length = [], [], 0
        for window in self._iter_spacy_windows(text):
            doc = spacy_nlp(window)
            for sent in doc.sents:
                sent_text = sent.text
                if current_length + len(sent_text) > self.chunk_size and current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk  = [current_chunk[-1]] if self.chunk_overlap > 0 else []
                    current_length = len(current_chunk[0]) if current_chunk else 0
                current_chunk.append(sent_text)
                current_length += len(sent_text)
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        return chunks

    def process_document(
        self,
        segments: List[Dict[str, Any]],
        source_file: str,
        user_email: str,
        progress: Optional[Callable[[float, str], None]] = None,
    ) -> str:
        """Chunk, embed, and upsert a document.

        `segments` is `[{'text': str, 'page': Optional[int]}, ...]`. For
        PDFs there is one segment per page; for every other file type
        there is exactly one segment with `page=None`. Chunking happens
        within each segment so a chunk's page attribution is unambiguous.
        """
        def report(frac: float, desc: str):
            if progress is not None:
                try:
                    progress(frac, desc=desc)
                except TypeError:
                    progress(frac)

        full_text = "\n\n".join(s["text"] for s in segments)
        content_hash = hashlib.sha256(f"{full_text}{user_email}".encode()).hexdigest()
        org = users[user_email]["org"] if user_email != "admin@example.com" else None

        report(0.05, "Checking for duplicate content...")
        existing_points = self.client.scroll(
            collection_name=self.collection_name,
            scroll_filter=Filter(must=[
                FieldCondition(key="user_email",    match=MatchValue(value=user_email)),
                FieldCondition(key="content_hash",  match=MatchValue(value=content_hash))
            ]),
            limit=1,
            with_payload=True
        )[0]

        if existing_points:
            existing_file = existing_points[0].payload.get("source_file", "unknown")
            return f"Duplicate content detected — already exists as '{existing_file}'. Upload skipped."

        report(0.10, f"Splitting {len(full_text):,} chars into chunks...")
        # Chunk each segment independently and remember which page each
        # resulting chunk came from. For PDFs this means no chunk ever
        # spans a page boundary; for other file types `page` stays None.
        chunked: List[Tuple[str, Optional[int]]] = []
        for seg in segments:
            seg_text = seg.get("text", "")
            seg_page = seg.get("page")
            if not seg_text.strip():
                continue
            for chunk in self.splitter(seg_text):
                chunked.append((chunk, seg_page))

        if not chunked:
            return f"No chunks generated for '{source_file}'. Upload failed."

        # Reserve a contiguous ID range up front. Holding the lock only for
        # this tiny critical section means two concurrent uploads each get
        # their own non-overlapping slice and never have to coordinate again.
        with self._id_lock:
            start_id = self.next_id
            self.next_id += len(chunked)

        # ── Embed batches in parallel ──────────────────────────────────────────
        batches: List[List[Tuple[str, Optional[int]]]] = [
            chunked[i:i + self.batch_size]
            for i in range(0, len(chunked), self.batch_size)
        ]
        total_batches = len(batches)
        report(0.20, f"Embedding {len(chunked):,} chunks "
                     f"in {total_batches} batches (×{min(EMBED_WORKERS, total_batches)} parallel)...")

        def _embed_batch(batch: List[Tuple[str, Optional[int]]]) -> List[List[float]]:
            return self.embeddings.embed_documents([c for c, _ in batch])

        workers = min(EMBED_WORKERS, total_batches)
        with ThreadPoolExecutor(max_workers=workers) as executor:
            # executor.map preserves order, which we need so payloads line up.
            all_embeds = list(executor.map(_embed_batch, batches))

        # ── Stream upserts batch-by-batch ──────────────────────────────────────
        # Smaller, incremental upserts keep memory low for huge PDFs and let
        # Qdrant make data searchable as soon as the first batch lands.
        report(0.70, "Uploading embeddings to Qdrant...")
        uploaded = 0
        point_offset = 0
        for batch_idx, (batch, embeds) in enumerate(zip(batches, all_embeds)):
            points = []
            for (chunk_text, chunk_page), emb in zip(batch, embeds):
                payload = {
                    "text":         chunk_text,
                    "source_file":  source_file,
                    "user_email":   user_email,
                    "content_hash": content_hash,
                    "org":          org,
                }
                # Only stamp `page` when it's meaningful (PDFs). Other file
                # types stay schema-compatible with older rows that pre-date
                # the page-tracking feature.
                if chunk_page is not None:
                    payload["page"] = chunk_page
                points.append(PointStruct(
                    id=start_id + point_offset,
                    vector=emb,
                    payload=payload,
                ))
                point_offset += 1
            if points:
                self.client.upsert(collection_name=self.collection_name, points=points)
                uploaded += len(points)
            frac = 0.70 + 0.30 * ((batch_idx + 1) / total_batches)
            report(frac, f"Uploaded {uploaded:,}/{len(chunked):,} chunks...")

        self.processed_files[source_file] = user_email
        return f"✅ Successfully processed '{source_file}' — {len(chunked)} chunks uploaded."

    def delete_by_source_file(self, source_file: str) -> int:
        try:
            f = Filter(must=[FieldCondition(key="source_file", match=MatchValue(value=source_file))])
            self.client.delete(collection_name=self.collection_name, points_selector=f)
            remaining = self.client.count(collection_name=self.collection_name,
                                          exact=False, count_filter=f).count
            if source_file in self.processed_files:
                del self.processed_files[source_file]
            return remaining
        except Exception as e:
            print(f"Error deleting {source_file}: {e}")
            return -1

    def search(self, query_vector: List[float], user_email: str, limit: int = 5):
        if user_email == "admin@example.com":
            results = self.client.query_points(
                collection_name=self.collection_name,
                query=query_vector, limit=limit, with_payload=True
            )
        else:
            org = users[user_email]["org"]
            results = self.client.query_points(
                collection_name=self.collection_name,
                query=query_vector,
                query_filter=Filter(must=[FieldCondition(key="org", match=MatchValue(value=org))]),
                limit=limit, with_payload=True
            )
        return results.points

    def get_processed_files(self, user_email: str = None):
        if user_email == "admin@example.com":
            return list(self.processed_files.keys())
        elif user_email:
            org = users[user_email]["org"]
            return [f for f, e in self.processed_files.items() if users[e]["org"] == org]
        return list(self.processed_files.keys())


# ══════════════════════════════════════════════════════════════════════════════
#  APP INIT
# ══════════════════════════════════════════════════════════════════════════════


processor = DocumentProcessor(
    collection_name="aiml_vector_db",
    chunk_strategy="semantic",
    chunk_size=800,
    chunk_overlap=150,
    embedding_model="text-embedding-3-small",
    batch_size=32
)

llm = ChatOpenAI(
    model=OPENAI_LLM_MODEL,
    temperature=0.2,
    openai_api_key=openai_api_key,
    timeout=60.0,
    max_retries=2,
)

# Raw OpenAI client needed for tool-use API
from openai import OpenAI
openai_client = OpenAI(api_key=openai_api_key, timeout=60.0, max_retries=2)


# ══════════════════════════════════════════════════════════════════════════════
#  GRADIO HANDLER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf", ".docx", ".doc",
    # Spreadsheets
    ".xlsx", ".xls",
    # Images
    ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp",
    # Text
    ".txt", ".md", ".csv", ".json", ".xml", ".html",
}


def process_file(file, current_files, selected_user_email,
                 progress=gr.Progress()):
    """
    Main upload handler.
    1. Validates file type.
    2. Runs the File-Type Agent to extract text.
    3. Passes extracted text to DocumentProcessor → Qdrant.
    """
    if file is None:
        return "No file uploaded.", "", gr.update(choices=current_files, value=[]), current_files

    filename  = os.path.basename(file.name)
    ext       = os.path.splitext(filename)[1].lower()
    user_email = selected_user_email

    if ext not in SUPPORTED_EXTENSIONS:
        msg = (f"❌ Unsupported file type: '{ext}'\n"
               f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        return msg, "", gr.update(choices=current_files, value=[]), current_files

    # Server-side size check. Gradio also enforces this at the upload boundary
    # (see `max_file_size` on demo.launch), but we re-validate here so a
    # client that bypasses the UI hint still gets a clean error message.
    try:
        size_bytes = os.path.getsize(file.name)
    except OSError:
        size_bytes = 0
    if size_bytes > MAX_UPLOAD_BYTES:
        size_mb = size_bytes / (1024 * 1024)
        msg = (f"❌ '{filename}' is {size_mb:.1f} MB, which exceeds the "
               f"{MAX_UPLOAD_MB} MB upload limit. Please split the file or "
               f"upload a smaller version.")
        return msg, "", gr.update(choices=current_files, value=[]), current_files

    try:
        progress(0.0, desc=f"Reading '{filename}' with the File-Type Agent...")
        extracted_text, agent_log, segments = file_type_agent(
            file.name, filename, openai_client
        )

        if not extracted_text.strip():
            return (f"⚠️ No text could be extracted from '{filename}'.",
                    agent_log,
                    gr.update(choices=current_files, value=[]),
                    current_files)

        status_message = processor.process_document(
            segments, filename, user_email, progress=progress
        )

        if "Successfully" in status_message:
            updated_files = processor.get_processed_files(user_email)
            return (status_message,
                    agent_log,
                    gr.update(choices=updated_files, value=[]),
                    updated_files)
        else:
            return (status_message,
                    agent_log,
                    gr.update(choices=current_files, value=[]),
                    current_files)

    except Exception as e:
        return (f"❌ Error processing '{filename}': {str(e)}",
                "",
                gr.update(choices=current_files, value=[]),
                current_files)


def delete_pdfs(filenames, current_files, selected_user_email):
    if not filenames:
        return "No files selected for deletion.", gr.update(choices=current_files, value=[]), current_files

    user_email   = selected_user_email
    deleted, errors = [], []

    for filename in filenames:
        if filename not in processor.processed_files or (
            processor.processed_files[filename] != user_email and user_email != "admin@example.com"
        ):
            errors.append(f"'{filename}' not found or not owned by {user_email}.")
            continue
        remaining = processor.delete_by_source_file(filename)
        if remaining == -1:
            errors.append(f"Failed to delete embeddings for '{filename}'.")
        elif remaining > 0:
            errors.append(f"Some embeddings for '{filename}' were not deleted ({remaining} remain).")
        else:
            deleted.append(filename)

    status = ""
    if deleted:
        status += f"✅ Deleted {len(deleted)} file(s): {', '.join(deleted)}."
    if errors:
        status += "\nErrors:\n" + "\n".join(errors)

    updated_files = processor.get_processed_files(user_email)
    return status, gr.update(choices=updated_files, value=[]), updated_files


def search_qdrant(query, selected_user_email):
    """Retrieve relevant passages and return them prefixed with source
    metadata. Returns `(context_str, source_pages)` where `source_pages`
    is a dict mapping `filename -> set[int]` of page numbers used (empty
    set when the file has no page info, e.g. for Word/Excel/text/images)."""
    try:
        query_embedding = processor.embeddings.embed_query(query)
        results = processor.search(query_vector=query_embedding, user_email=selected_user_email, limit=5)

        if not results:
            return "No relevant information found.", {}

        # Step 1: find the best score per file
        file_best_score = {}
        for result in results:
            sf = result.payload.get("source_file", "Unknown")
            file_best_score[sf] = max(file_best_score.get(sf, 0), result.score)

        # Step 2: only keep files whose best score is within 0.15 of the top file
        # e.g. top file scores 0.82, threshold = 0.67 — car insurance at 0.55 gets dropped
        top_score = max(file_best_score.values())
        relevant_files = {f for f, s in file_best_score.items() if s >= top_score - 0.15}

        # Build context with per-chunk source headers so the LLM can attribute
        # specific claims to specific pages when answering.
        context_parts: List[str] = []
        source_pages: Dict[str, set] = {}
        for result in results:
            source_file = result.payload.get("source_file", "Unknown")
            if source_file not in relevant_files:
                continue
            text = result.payload.get("text", "")
            page = result.payload.get("page")
            header = f"[Source: {source_file}"
            if page is not None:
                header += f", page {page}"
            header += "]"
            context_parts.append(f"{header}\n{text}")
            page_set = source_pages.setdefault(source_file, set())
            if page is not None:
                page_set.add(page)

        context = "\n\n".join(context_parts)
        return context or "No relevant information found.", source_pages
    except Exception as e:
        return f"Error searching Qdrant: {str(e)}", {}


def _format_sources_line(source_pages: Dict[str, set]) -> str:
    """Render the trailing `Sources:` line, including per-file page numbers
    when available. Example:
        Sources: report.pdf (pages 3, 4, 7), notes.docx
    """
    if not source_pages:
        return "Sources: None identified"
    parts = []
    for filename in sorted(source_pages.keys()):
        pages = sorted(source_pages[filename])
        if pages:
            label = "page" if len(pages) == 1 else "pages"
            page_str = ", ".join(str(p) for p in pages)
            parts.append(f"{filename} ({label} {page_str})")
        else:
            parts.append(filename)
    return "Sources: " + ", ".join(parts)


# Matches a "Sources:" / "Source:" / "**Sources:**" / "### Sources" line in
# the LLM's response and everything that follows it. We anchor with `\Z`
# (end of string) so `.*` (DOTALL) gobbles all of the model's own attribution
# block — header line plus any bullets or follow-up lines — letting us
# append our single canonical Sources line cleanly.
#
# Earlier version required the header line to end immediately after the
# colon (`(?=\n|$)`), which silently failed on the common case
# `Sources: file.pdf (pages 1, 2, 3)` because the colon is followed by the
# citation list, not a newline. The new pattern uses `\b` after `sources?`
# so it still won't fire on words like "sourcecode" or "outsourced", but
# doesn't care what comes after on the same line.
_TRAILING_SOURCES_RE = re.compile(
    r"\n[ \t]*(?:[#>\-*][ \t]*)*\**[ \t]*sources?\b.*\Z",
    re.IGNORECASE | re.DOTALL,
)


def _strip_trailing_sources(text: str) -> str:
    """Remove any trailing `Sources:` block from the LLM's answer.

    GPT-3.5 sometimes appends its own attribution list at the end of the
    response despite the system-prompt instruction not to. If we don't
    strip it, the user sees two `Sources:` lines (one from the model, one
    we control), which looks like a bug. The regex matches from the last
    `Sources:` heading to the end of the string."""
    return _TRAILING_SOURCES_RE.sub("", text).rstrip()


def chatbot_response(message, history, selected_user_email):
    try:
        context, source_pages = search_qdrant(message, selected_user_email)
        system_prompt = (
            "You are a helpful assistant that answers questions based on the "
            "provided document context. Each excerpt below is prefixed with a "
            "header like `[Source: <file>, page <N>]` indicating where it came "
            "from. When you state a fact taken from a specific excerpt, cite "
            "the page number inline in parentheses, e.g. `(page 3)`. If the "
            "excerpt has no page (non-PDF source), cite the filename instead. "
            "Use only the context below. If it does not contain the answer, "
            "say so explicitly rather than guessing.\n\n"
            "VERY IMPORTANT: do NOT add a 'Sources:' section, a 'References:' "
            "section, or any kind of source list at the end of your answer. "
            "A canonical Sources line is appended automatically after your "
            "response; if you add one yourself the user will see two. Inline "
            "`(page N)` citations inside sentences are fine and encouraged.\n\n"
            f"Context:\n{context}"
        )
        messages = [{"role": "system", "content": system_prompt}]
        for msg in history:
            role = "user" if msg["role"] == "user" else "assistant"
            messages.append({"role": role, "content": msg["content"]})
        messages.append({"role": "user", "content": message})

        response = llm.invoke(messages).content

        # Belt-and-braces: even with the instruction above, GPT-3.5 sometimes
        # still emits a trailing "Sources:" block. Strip it before appending
        # our canonical one so the user only ever sees a single Sources line.
        response = _strip_trailing_sources(response)
        response += "\n\n" + _format_sources_line(source_pages)
        return response
    except Exception as e:
        return f"Error: {str(e)}\n\nSources: None identified"


def chat_handler(message, history, selected_user_email):
    # Gradio Chatbot expects a list of messages in the format:
    #   [{"role": "user", "content": "..."}, {"role": "assistant", "content": "..."}, ...]
    # Ensure compatibility whether Gradio passes tuples or dicts.
    if history is None:
        history = []

    normalized_history = []
    for pair in history:
        if isinstance(pair, (list, tuple)) and len(pair) == 2:
            normalized_history.append({"role": "user", "content": pair[0] or ""})
            normalized_history.append({"role": "assistant", "content": pair[1] or ""})
        elif isinstance(pair, dict) and "role" in pair and "content" in pair:
            normalized_history.append(pair)

    response = chatbot_response(message, normalized_history, selected_user_email)
    normalized_history.append({"role": "user", "content": message})
    normalized_history.append({"role": "assistant", "content": response})
    return normalized_history, ""


def clear_chat():
    return []


def update_file_list(selected_user_email):
    files = processor.get_processed_files(selected_user_email)
    return gr.update(choices=files, value=[])


def refresh_on_load(selected_user_email):
    # Re-read the latest state from Qdrant so a browser reload after a delete
    # doesn't show stale choices captured at app-startup time.
    processor.processed_files = processor._fetch_processed_files()
    files = processor.get_processed_files(selected_user_email)
    return gr.update(choices=files, value=[]), files


# ══════════════════════════════════════════════════════════════════════════════
#  GRADIO UI
# ══════════════════════════════════════════════════════════════════════════════
with gr.Blocks(title="Document Reader") as demo:
    gr.Markdown("# 📄 Document Reader with Qdrant")
    gr.Markdown(
        "Upload **PDF, Word, Excel, Images, or Text files**. "
        "An AI agent automatically picks the right reader for each file type."
    )

    file_state = gr.State(value=processor.get_processed_files())
    user_state = gr.State(value="john.doe123@example.com")

    with gr.Row():
        # ── LEFT: Chat ────────────────────────────────────────────────────────
        with gr.Column(scale=1):
            gr.Markdown("## 💬 Chat with Your Documents")
            chatbot = gr.Chatbot(label="Document Chatbot", height=400)
            msg     = gr.Textbox(
                label="Your Message",
                placeholder="Ask anything about your uploaded documents..."
            )
            with gr.Row():
                submit_btn = gr.Button("Submit", variant="primary")
                clear_btn  = gr.Button("Clear Chat")

            # Wire the same handler to both the Submit button click and the
            # Enter key in the textbox so the user has two equivalent ways
            # to send a message.
            chat_inputs  = [msg, chatbot, user_state]
            chat_outputs = [chatbot, msg]
            submit_btn.click(fn=chat_handler, inputs=chat_inputs, outputs=chat_outputs)
            msg.submit(fn=chat_handler,       inputs=chat_inputs, outputs=chat_outputs)
            clear_btn.click(fn=clear_chat, outputs=chatbot)

        # ── RIGHT: Upload & Manage ────────────────────────────────────────────
        with gr.Column(scale=1):
            with gr.Group():
                gr.Markdown("## 👤 Select User")
                user_dropdown = gr.Dropdown(
                    label="User",
                    choices=list(users.keys()),
                    value="john.doe123@example.com",
                    interactive=True
                )

            with gr.Group():
                gr.Markdown("## 📤 Upload File")
                gr.Markdown(
                    "_Supported: PDF · Word (.docx) · Excel (.xlsx) · "
                    "Images (png/jpg/jpeg/tiff/bmp/webp) · Text (txt/md/csv/json/xml/html)_  \n"
                    f"_Max upload size: **{MAX_UPLOAD_MB} MB**._"
                )
                file_input  = gr.File(
                    label="Upload File",
                    file_types=list(SUPPORTED_EXTENSIONS)
                )
                upload_btn    = gr.Button("Process File", variant="primary")
                upload_status = gr.Textbox(label="Upload Status", interactive=False)

                # Agent reasoning log — shows users what the agent decided
                with gr.Accordion("🤖 Agent Log (what the agent did)", open=False):
                    agent_log_box = gr.Textbox(
                        label="File-Type Agent Reasoning",
                        lines=6,
                        interactive=False
                    )

            with gr.Group():
                gr.Markdown("## 🗂️ Manage Documents")
                delete_checkboxes = gr.CheckboxGroup(
                    label="Select Files to Delete",
                    value=[],
                    choices=processor.get_processed_files("john.doe123@example.com"),
                    interactive=True
                )
                delete_btn    = gr.Button("Delete Selected Files", variant="stop")
                delete_status = gr.Textbox(label="Delete Status", interactive=False)

    # ── Event wiring ──────────────────────────────────────────────────────────
    user_dropdown.change(
        fn=update_file_list,
        inputs=[user_dropdown],
        outputs=[delete_checkboxes]
    ).then(
        fn=lambda x: x,
        inputs=[user_dropdown],
        outputs=[user_state]
    )

    upload_btn.click(
        fn=process_file,
        inputs=[file_input, file_state, user_state],
        outputs=[upload_status, agent_log_box, delete_checkboxes, file_state]
    )

    delete_btn.click(
        fn=delete_pdfs,
        inputs=[delete_checkboxes, file_state, user_state],
        outputs=[delete_status, delete_checkboxes, file_state]
    )

    # Re-sync the file list with Qdrant every time the page is (re)loaded.
    # Without this, gr.CheckboxGroup keeps the choices that were captured at
    # app startup, so deleted files keep reappearing after a browser refresh.
    demo.load(
        fn=refresh_on_load,
        inputs=[user_state],
        outputs=[delete_checkboxes, file_state]
    )

if __name__ == "__main__":
    # `queue()` is required for gr.Progress and lets long-running uploads run
    # past Gradio's default per-request timeout. Concurrency is kept modest so
    # a single beefy PDF can use the full embedding thread pool without
    # contending with other users on the same OpenAI quota.
    demo.queue(default_concurrency_limit=2, max_size=16).launch(
        max_file_size=f"{MAX_UPLOAD_MB}mb",
    )
